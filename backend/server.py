from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime
import aiosqlite
import json
import re
import math
from collections import Counter, defaultdict
import asyncio
from openai import AsyncAzureOpenAI

# Document processing imports
import PyPDF2
import docx
from bs4 import BeautifulSoup
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Azure OpenAI client setup
azure_client = AsyncAzureOpenAI(
    api_key=os.environ['AZURE_OPENAI_API_KEY'],
    api_version=os.environ['AZURE_OPENAI_API_VERSION'],
    azure_endpoint=os.environ['AZURE_OPENAI_ENDPOINT']
)

# Database path
DATABASE_PATH = ROOT_DIR / os.environ.get('DATABASE_PATH', 'deepdive_rag.db')

# Create the main app without a prefix
app = FastAPI(title="DeepDive RAG API", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Database Models
class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    content: str
    file_type: str
    upload_date: datetime = Field(default_factory=datetime.utcnow)
    chunk_count: int = 0

class DocumentChunk(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    content: str
    chunk_index: int
    word_count: int
    tf_idf_vector: Dict[str, float] = Field(default_factory=dict)

class ResearchSession(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    query: str
    status: str = "processing"  # processing, completed, failed
    created_at: datetime = Field(default_factory=datetime.utcnow)
    step1_chunks: List[str] = Field(default_factory=list)
    step1_answer: str = ""
    gap_questions: List[str] = Field(default_factory=list)
    step3_chunks: List[str] = Field(default_factory=list)
    final_answer: str = ""
    citations: List[Dict[str, Any]] = Field(default_factory=list)

class QueryRequest(BaseModel):
    query: str

class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    chunk_count: int
    message: str

class DocumentDeleteResponse(BaseModel):
    message: str
    deleted_document_id: str
    deleted_chunks: int

class ResearchResponse(BaseModel):
    session_id: str
    query: str
    status: str
    step1_answer: Optional[str] = None
    gap_questions: List[str] = Field(default_factory=list)
    final_answer: Optional[str] = None
    citations: List[Dict[str, Any]] = Field(default_factory=list)
    timeline: List[Dict[str, Any]] = Field(default_factory=list)

# Database initialization
async def init_database():
    """Initialize SQLite database with required tables"""
    async with aiosqlite.connect(DATABASE_PATH) as db:
        # Documents table
        await db.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                content TEXT NOT NULL,
                file_type TEXT NOT NULL,
                upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                chunk_count INTEGER DEFAULT 0
            )
        ''')
        
        # Document chunks table
        await db.execute('''
            CREATE TABLE IF NOT EXISTS document_chunks (
                id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL,
                content TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                word_count INTEGER NOT NULL,
                tf_idf_vector TEXT,
                FOREIGN KEY (document_id) REFERENCES documents (id)
            )
        ''')
        
        # Research sessions table
        await db.execute('''
            CREATE TABLE IF NOT EXISTS research_sessions (
                id TEXT PRIMARY KEY,
                query TEXT NOT NULL,
                status TEXT DEFAULT 'processing',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                step1_chunks TEXT,
                step1_answer TEXT,
                gap_questions TEXT,
                step3_chunks TEXT,
                final_answer TEXT,
                citations TEXT
            )
        ''')
        
        await db.commit()

class DocumentProcessor:
    """Handles document processing and chunking"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_html(file_content: bytes) -> str:
        """Extract text from HTML file"""
        try:
            soup = BeautifulSoup(file_content, 'html.parser')
            return soup.get_text()
        except Exception as e:
            logging.error(f"Error extracting text from HTML: {e}")
            return ""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk.strip())
        
        return chunks

class RetrievalEngine:
    """Handles document retrieval using TF-IDF and cosine similarity"""
    
    @staticmethod
    def preprocess_text(text: str) -> List[str]:
        """Basic text preprocessing"""
        # Convert to lowercase and remove special characters
        text = re.sub(r'[^a-zA-Z0-9\s]', '', text.lower())
        words = text.split()
        # Remove common stop words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should'}
        return [word for word in words if word not in stop_words and len(word) > 2]
    
    @staticmethod
    def calculate_tf_idf(chunks: List[str]) -> Dict[str, Dict[str, float]]:
        """Calculate TF-IDF vectors for all chunks"""
        # Preprocess all chunks
        processed_chunks = [RetrievalEngine.preprocess_text(chunk) for chunk in chunks]
        
        # Calculate document frequencies
        all_words = set()
        for chunk_words in processed_chunks:
            all_words.update(chunk_words)
        
        document_frequencies = {}
        for word in all_words:
            df = sum(1 for chunk_words in processed_chunks if word in chunk_words)
            document_frequencies[word] = df
        
        # Calculate TF-IDF for each chunk
        tfidf_vectors = {}
        total_docs = len(chunks)
        
        for i, chunk_words in enumerate(processed_chunks):
            word_counts = Counter(chunk_words)
            total_words = len(chunk_words)
            
            tfidf_vector = {}
            for word in word_counts:
                tf = word_counts[word] / total_words
                idf = math.log(total_docs / document_frequencies[word])
                tfidf_vector[word] = tf * idf
            
            tfidf_vectors[str(i)] = tfidf_vector
        
        return tfidf_vectors
    
    @staticmethod
    def cosine_similarity(vec1: Dict[str, float], vec2: Dict[str, float]) -> float:
        """Calculate cosine similarity between two TF-IDF vectors"""
        # Get all unique words
        all_words = set(vec1.keys()) | set(vec2.keys())
        
        # Calculate dot product and magnitudes
        dot_product = sum(vec1.get(word, 0) * vec2.get(word, 0) for word in all_words)
        magnitude1 = math.sqrt(sum(val ** 2 for val in vec1.values()))
        magnitude2 = math.sqrt(sum(val ** 2 for val in vec2.values()))
        
        if magnitude1 == 0 or magnitude2 == 0:
            return 0
        
        return dot_product / (magnitude1 * magnitude2)
    
    @staticmethod
    async def search_chunks(query: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """Search for relevant chunks using hybrid approach"""
        async with aiosqlite.connect(DATABASE_PATH) as db:
            # Get all chunks
            cursor = await db.execute('''
                SELECT dc.id, dc.content, dc.document_id, dc.chunk_index, dc.tf_idf_vector, d.filename
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
            ''')
            chunks = await cursor.fetchall()
            
            if not chunks:
                return []
            
            # Preprocess query
            query_words = RetrievalEngine.preprocess_text(query)
            query_word_counts = Counter(query_words)
            
            # Calculate query TF-IDF vector
            query_tfidf = {}
            total_query_words = len(query_words)
            for word in query_word_counts:
                tf = query_word_counts[word] / total_query_words
                # For query, we'll use a simple IDF approximation
                query_tfidf[word] = tf
            
            # Score each chunk
            scored_chunks = []
            for chunk in chunks:
                chunk_id, content, doc_id, chunk_index, tfidf_json, filename = chunk
                
                # Parse TF-IDF vector
                try:
                    chunk_tfidf = json.loads(tfidf_json) if tfidf_json else {}
                except:
                    chunk_tfidf = {}
                
                # Calculate semantic similarity
                semantic_score = RetrievalEngine.cosine_similarity(query_tfidf, chunk_tfidf)
                
                # Calculate keyword match score
                chunk_words = RetrievalEngine.preprocess_text(content)
                keyword_matches = sum(1 for word in query_words if word in chunk_words)
                keyword_score = keyword_matches / len(query_words) if query_words else 0
                
                # Combine scores (weighted)
                final_score = 0.6 * semantic_score + 0.4 * keyword_score
                
                scored_chunks.append({
                    'id': chunk_id,
                    'content': content,
                    'document_id': doc_id,
                    'chunk_index': chunk_index,
                    'filename': filename,
                    'score': final_score,
                    'semantic_score': semantic_score,
                    'keyword_score': keyword_score
                })
            
            # Sort by score and return top k
            scored_chunks.sort(key=lambda x: x['score'], reverse=True)
            return scored_chunks[:top_k]

class RAGPipeline:
    """Main RAG pipeline orchestrator"""
    
    @staticmethod
    async def process_query(query: str) -> str:
        """Execute the complete 5-step RAG pipeline"""
        session_id = str(uuid.uuid4())
        
        try:
            # Step 1: Initial Retrieval
            step1_chunks = await RetrievalEngine.search_chunks(query, top_k=8)
            
            if not step1_chunks:
                return session_id, "No documents found. Please upload some documents first."
            
            # Step 2: First Generation
            context = "\n\n".join([f"Source {i+1}: {chunk['content']}" for i, chunk in enumerate(step1_chunks)])
            
            first_generation_prompt = f"""
            Query: {query}
            
            Context from retrieved documents:
            {context}
            
            Based on the provided context, provide an initial answer to the query. Then, identify any gaps or unclear points that need more information to provide a complete answer.
            
            Format your response as:
            INITIAL_ANSWER: [Your initial answer here]
            
            GAP_QUESTIONS: [List specific questions or topics that need more information, separated by semicolons]
            """
            
            response = await azure_client.chat.completions.create(
                model=os.environ['AZURE_OPENAI_DEPLOYMENT_NAME'],
                messages=[{"role": "user", "content": first_generation_prompt}],
                temperature=0.7,
                max_tokens=1000
            )
            
            first_response = response.choices[0].message.content
            
            # Parse initial answer and gap questions
            initial_answer = ""
            gap_questions = []
            
            if "INITIAL_ANSWER:" in first_response:
                parts = first_response.split("GAP_QUESTIONS:")
                initial_answer = parts[0].replace("INITIAL_ANSWER:", "").strip()
                if len(parts) > 1:
                    gap_text = parts[1].strip()
                    gap_questions = [q.strip() for q in gap_text.split(";") if q.strip()]
            
            # Step 3: Secondary Retrieval
            step3_chunks = []
            if gap_questions:
                for gap_question in gap_questions:
                    additional_chunks = await RetrievalEngine.search_chunks(gap_question, top_k=5)
                    step3_chunks.extend(additional_chunks)
            
            # Remove duplicates and combine chunks
            all_chunks = step1_chunks + step3_chunks
            unique_chunks = []
            seen_ids = set()
            for chunk in all_chunks:
                if chunk['id'] not in seen_ids:
                    unique_chunks.append(chunk)
                    seen_ids.add(chunk['id'])
            
            # Step 4: Final Generation
            final_context = "\n\n".join([f"Source {i+1} ({chunk['filename']}): {chunk['content']}" for i, chunk in enumerate(unique_chunks)])
            
            final_generation_prompt = f"""
            Original Query: {query}
            
            Initial Answer: {initial_answer}
            
            Additional Context from comprehensive search:
            {final_context}
            
            Based on the original query, initial answer, and comprehensive context, provide a refined, complete answer with proper citations.
            
            Format your response as:
            FINAL_ANSWER: [Your comprehensive final answer with [Source X] citations throughout]
            
            Use [Source X] format to cite sources, where X corresponds to the source number in the context above.
            """
            
            final_response = await azure_client.chat.completions.create(
                model=os.environ['AZURE_OPENAI_DEPLOYMENT_NAME'],
                messages=[{"role": "user", "content": final_generation_prompt}],
                temperature=0.7,
                max_tokens=1500
            )
            
            final_answer = final_response.choices[0].message.content
            if "FINAL_ANSWER:" in final_answer:
                final_answer = final_answer.replace("FINAL_ANSWER:", "").strip()
            
            # Create citations
            citations = []
            for i, chunk in enumerate(unique_chunks):
                citations.append({
                    'id': chunk['id'],
                    'source_number': i + 1,
                    'filename': chunk['filename'],
                    'content': chunk['content'][:200] + "..." if len(chunk['content']) > 200 else chunk['content'],
                    'score': chunk['score']
                })
            
            # Save session to database
            async with aiosqlite.connect(DATABASE_PATH) as db:
                await db.execute('''
                    INSERT INTO research_sessions 
                    (id, query, status, step1_chunks, step1_answer, gap_questions, step3_chunks, final_answer, citations)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    session_id,
                    query,
                    "completed",
                    json.dumps([chunk['id'] for chunk in step1_chunks]),
                    initial_answer,
                    json.dumps(gap_questions),
                    json.dumps([chunk['id'] for chunk in step3_chunks]),
                    final_answer,
                    json.dumps(citations)
                ))
                await db.commit()
            
            return session_id, {
                'session_id': session_id,
                'query': query,
                'status': 'completed',
                'step1_answer': initial_answer,
                'gap_questions': gap_questions,
                'final_answer': final_answer,
                'citations': citations,
                'timeline': [
                    {'step': 1, 'description': f'Retrieved {len(step1_chunks)} initial chunks', 'chunks': len(step1_chunks)},
                    {'step': 2, 'description': f'Generated initial answer and identified {len(gap_questions)} gaps', 'gaps': len(gap_questions)},
                    {'step': 3, 'description': f'Retrieved {len(step3_chunks)} additional chunks for gaps', 'chunks': len(step3_chunks)},
                    {'step': 4, 'description': 'Generated final comprehensive answer with citations', 'citations': len(citations)},
                    {'step': 5, 'description': 'Ready for display', 'total_sources': len(unique_chunks)}
                ]
            }
            
        except Exception as e:
            logging.error(f"Error in RAG pipeline: {e}")
            # Save failed session
            async with aiosqlite.connect(DATABASE_PATH) as db:
                await db.execute('''
                    INSERT INTO research_sessions (id, query, status)
                    VALUES (?, ?, ?)
                ''', (session_id, query, "failed"))
                await db.commit()
            
            return session_id, f"Error processing query: {str(e)}"

# API Routes
@api_router.get("/")
async def root():
    return {"message": "DeepDive RAG API is running"}

@api_router.post("/upload-document", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a document"""
    try:
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        filename = file.filename
        file_type = filename.split('.')[-1].lower()
        
        if file_type == 'pdf':
            text = DocumentProcessor.extract_text_from_pdf(content)
        elif file_type == 'docx':
            text = DocumentProcessor.extract_text_from_docx(content)
        elif file_type in ['html', 'htm']:
            text = DocumentProcessor.extract_text_from_html(content)
        elif file_type == 'txt':
            text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, TXT, or HTML files.")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in file")
        
        # Create document
        document_id = str(uuid.uuid4())
        
        # Chunk the text
        chunks = DocumentProcessor.chunk_text(text)
        
        # Calculate TF-IDF vectors for chunks
        tfidf_vectors = RetrievalEngine.calculate_tf_idf(chunks)
        
        # Save to database
        async with aiosqlite.connect(DATABASE_PATH) as db:
            # Save document
            await db.execute('''
                INSERT INTO documents (id, filename, content, file_type, chunk_count)
                VALUES (?, ?, ?, ?, ?)
            ''', (document_id, filename, text, file_type, len(chunks)))
            
            # Save chunks
            for i, chunk in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                word_count = len(chunk.split())
                tfidf_json = json.dumps(tfidf_vectors.get(str(i), {}))
                
                await db.execute('''
                    INSERT INTO document_chunks 
                    (id, document_id, content, chunk_index, word_count, tf_idf_vector)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (chunk_id, document_id, chunk, i, word_count, tfidf_json))
            
            await db.commit()
        
        return DocumentUploadResponse(
            document_id=document_id,
            filename=filename,
            chunk_count=len(chunks),
            message=f"Document '{filename}' processed successfully into {len(chunks)} chunks"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@api_router.delete("/documents/{document_id}", response_model=DocumentDeleteResponse)
async def delete_document(document_id: str):
    """Delete a document and all its associated chunks"""
    try:
        async with aiosqlite.connect(DATABASE_PATH) as db:
            # First, check if document exists
            cursor = await db.execute('''
                SELECT filename, chunk_count FROM documents WHERE id = ?
            ''', (document_id,))
            document = await cursor.fetchone()
            
            if not document:
                raise HTTPException(status_code=404, detail="Document not found")
            
            filename, chunk_count = document
            
            # Delete all chunks associated with this document
            await db.execute('''
                DELETE FROM document_chunks WHERE document_id = ?
            ''', (document_id,))
            
            # Delete the document
            await db.execute('''
                DELETE FROM documents WHERE id = ?
            ''', (document_id,))
            
            await db.commit()
            
            return DocumentDeleteResponse(
                message=f"Document '{filename}' and {chunk_count} associated chunks deleted successfully",
                deleted_document_id=document_id,
                deleted_chunks=chunk_count
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@api_router.delete("/documents")
async def delete_all_documents():
    """Delete all documents and chunks"""
    try:
        async with aiosqlite.connect(DATABASE_PATH) as db:
            # Count documents before deletion
            cursor = await db.execute('SELECT COUNT(*) FROM documents')
            doc_count = (await cursor.fetchone())[0]
            
            cursor = await db.execute('SELECT COUNT(*) FROM document_chunks')
            chunk_count = (await cursor.fetchone())[0]
            
            # Delete all chunks first (due to foreign key constraint)
            await db.execute('DELETE FROM document_chunks')
            
            # Delete all documents
            await db.execute('DELETE FROM documents')
            
            # Delete all research sessions
            await db.execute('DELETE FROM research_sessions')
            
            await db.commit()
            
            return {
                "message": f"All documents deleted successfully",
                "deleted_documents": doc_count,
                "deleted_chunks": chunk_count
            }
            
    except Exception as e:
        logging.error(f"Error deleting all documents: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting documents: {str(e)}")

@api_router.post("/research", response_model=ResearchResponse)
async def research_query(request: QueryRequest):
    """Process a research query through the RAG pipeline"""
    try:
        if not request.query or not request.query.strip():
            raise HTTPException(status_code=400, detail="Query cannot be empty")
            
        session_id, result = await RAGPipeline.process_query(request.query.strip())
        
        if isinstance(result, str):
            # Error occurred
            return ResearchResponse(
                session_id=session_id,
                query=request.query,
                status="failed",
                final_answer=result
            )
        
        return ResearchResponse(**result)
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error processing research query: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing research: {str(e)}")

@api_router.get("/research/{session_id}", response_model=ResearchResponse)
async def get_research_session(session_id: str):
    """Get research session details"""
    try:
        async with aiosqlite.connect(DATABASE_PATH) as db:
            cursor = await db.execute('''
                SELECT * FROM research_sessions WHERE id = ?
            ''', (session_id,))
            session = await cursor.fetchone()
            
            if not session:
                raise HTTPException(status_code=404, detail="Research session not found")
            
            # Parse JSON fields
            citations = json.loads(session[9]) if session[9] else []
            gap_questions = json.loads(session[6]) if session[6] else []
            
            return ResearchResponse(
                session_id=session[0],
                query=session[1],
                status=session[2],
                step1_answer=session[4],
                gap_questions=gap_questions,
                final_answer=session[8],
                citations=citations
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error getting research session: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving session: {str(e)}")

@api_router.get("/documents")
async def list_documents():
    """List all uploaded documents"""
    try:
        async with aiosqlite.connect(DATABASE_PATH) as db:
            cursor = await db.execute('''
                SELECT id, filename, file_type, upload_date, chunk_count 
                FROM documents ORDER BY upload_date DESC
            ''')
            documents = await cursor.fetchall()
            
            return [
                {
                    'id': doc[0],
                    'filename': doc[1],
                    'file_type': doc[2],
                    'upload_date': doc[3],
                    'chunk_count': doc[4]
                }
                for doc in documents
            ]
            
    except Exception as e:
        logging.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving documents: {str(e)}")

@api_router.get("/documents/{document_id}")
async def get_document_details(document_id: str):
    """Get detailed information about a specific document"""
    try:
        async with aiosqlite.connect(DATABASE_PATH) as db:
            # Get document details
            cursor = await db.execute('''
                SELECT id, filename, file_type, upload_date, chunk_count, content
                FROM documents WHERE id = ?
            ''', (document_id,))
            document = await cursor.fetchone()
            
            if not document:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Get chunks for this document
            cursor = await db.execute('''
                SELECT id, content, chunk_index, word_count
                FROM document_chunks WHERE document_id = ?
                ORDER BY chunk_index
            ''', (document_id,))
            chunks = await cursor.fetchall()
            
            return {
                'id': document[0],
                'filename': document[1],
                'file_type': document[2],
                'upload_date': document[3],
                'chunk_count': document[4],
                'content_preview': document[5][:500] + "..." if len(document[5]) > 500 else document[5],
                'chunks': [
                    {
                        'id': chunk[0],
                        'content': chunk[1],
                        'chunk_index': chunk[2],
                        'word_count': chunk[3]
                    }
                    for chunk in chunks
                ]
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error getting document details: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving document details: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup_event():
    """Initialize database on startup"""
    await init_database()
    logger.info("DeepDive RAG API started successfully")

@app.on_event("shutdown")
async def shutdown_event():
    """Cleanup on shutdown"""
    logger.info("DeepDive RAG API shutting down")